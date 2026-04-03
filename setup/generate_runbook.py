#!/usr/bin/env python3
"""Generate a runbook with screenshots from the orchestrator UI.

Uses Playwright (headless Chrome) to navigate the UI, capture screenshots,
and generates a Word document with embedded images and step-by-step instructions.

Prerequisites:
    pip install playwright python-docx
    playwright install chromium

Usage:
    python generate_runbook.py --url http://s4ll2tstapp0013:8000 --user admin --password admin
    python generate_runbook.py --url http://localhost:8000 --output runbook.docx
"""

import argparse
import os
import sys
import time
from datetime import datetime
from pathlib import Path

try:
    from playwright.sync_api import sync_playwright
except ImportError:
    print("Install playwright: pip install playwright && playwright install chromium")
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("Install python-docx: pip install python-docx")
    sys.exit(1)


class RunbookGenerator:
    def __init__(self, base_url, username, password, output_dir):
        self.base_url = base_url.rstrip("/")
        self.username = username
        self.password = password
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.screenshots = []  # list of (filename, title, description)
        self.page = None
        self.browser = None

    def _screenshot(self, name, title, description=""):
        """Take a screenshot and record metadata."""
        path = str(self.output_dir / f"{name}.png")
        self.page.screenshot(path=path, full_page=True)
        self.screenshots.append((path, title, description))
        print(f"  [SNAP] {name}: {title}")
        return path

    def _wait_and_screenshot(self, name, title, description="", wait_ms=1000):
        """Wait for page to settle, then screenshot."""
        self.page.wait_for_load_state("networkidle")
        time.sleep(wait_ms / 1000)
        return self._screenshot(name, title, description)

    def _goto(self, path):
        """Navigate to a path relative to base_url."""
        url = f"{self.base_url}{path}"
        self.page.goto(url, wait_until="networkidle")
        time.sleep(1)

    def login(self):
        """Login to the orchestrator."""
        print("\n[LOGIN]")
        self._goto("/login")
        time.sleep(1)

        # Fill login form
        try:
            self.page.fill('input[name="username"], input[type="text"]', self.username)
            self.page.fill('input[name="password"], input[type="password"]', self.password)
            self._screenshot("01_login", "Login Page",
                             "Navigate to the orchestrator URL and enter your credentials.")

            # Click login button
            self.page.click('button[type="submit"], input[type="submit"], .btn-primary')
            self.page.wait_for_load_state("networkidle")
            time.sleep(2)
        except Exception as e:
            print(f"  Login form not found or already logged in: {e}")

    def capture_dashboard(self):
        """Capture the admin dashboard."""
        print("\n[DASHBOARD]")
        self._goto("/admin/dashboard")
        self._wait_and_screenshot("02_dashboard", "Admin Dashboard",
            "The main dashboard shows an overview of labs, servers, "
            "active and completed baseline tests, and recent test runs.")

    def capture_servers(self):
        """Capture the servers page."""
        print("\n[SERVERS]")
        self._goto("/admin/servers")
        self._wait_and_screenshot("03_servers", "Servers Configuration",
            "View all registered servers including targets, load generators, "
            "and their IP addresses, OS family, and lab assignment.")

    def capture_load_profiles(self):
        """Capture load profiles page."""
        print("\n[LOAD PROFILES]")
        self._goto("/admin/load-profiles")
        self._wait_and_screenshot("04_load_profiles", "Load Profiles",
            "Load profiles define target CPU ranges for calibration. "
            "Typically: low (20-40%), medium (40-60%), high (60-80%).")

    def capture_agents(self):
        """Capture agents page."""
        print("\n[AGENTS]")
        self._goto("/admin/agents")
        self._wait_and_screenshot("05_agents", "Agents",
            "View registered security agents and their versions. "
            "Agents are installed on targets and their performance overhead is measured.")

    def capture_agent_sets(self):
        """Capture agent sets page."""
        print("\n[AGENT SETS]")
        self._goto("/admin/subgroup-definitions")
        self._wait_and_screenshot("06_agent_sets", "Agent Sets",
            "Agent Sets define combinations of security agents for testing.\n"
            "Examples: 'CrowdStrike Only' (single agent), 'CS + Tanium' (multi-agent).\n"
            "Agent Sets are reused across servers — define once, test everywhere.")

    def capture_snapshot_manager(self):
        """Capture snapshot manager."""
        print("\n[SNAPSHOT MANAGER]")
        self._goto("/snapshots")
        self._wait_and_screenshot("07_snapshots", "Snapshot Manager",
            "View the snapshot hierarchy per server.\n"
            "OS Level = root snapshot (clean prepared OS).\n"
            "Subgroups = agent sets with snapshots taken after agent installation.")

    def capture_tests_page(self):
        """Capture the unified Tests page (OS Level + Agent Set)."""
        print("\n[TESTS PAGE]")
        self._goto("/os-level-tests")
        self._wait_and_screenshot("08_tests_page", "Tests — OS Level & Agent Set",
            "The Tests page shows OS Level Tests (baselines) as expandable rows.\n"
            "Expand a row to see Agent Set Tests nested under it.\n"
            "- 'New OS Level Test': Create a clean OS baseline test.\n"
            "- 'Agent Set Test' (on completed baselines): Create a compare test with agents.")

    def capture_baseline_test_list(self):
        """Capture the baseline test list (All Tests)."""
        print("\n[ALL TESTS]")
        self._goto("/baseline-tests")
        self._wait_and_screenshot("09_all_tests", "All Baseline Tests",
            "Flat list of all test runs (both OS Level and Agent Set).\n"
            "Shows type, state, verdict, and action buttons.\n"
            "Use 'Delete All Tests' to clean up.")

    def capture_create_new_baseline(self):
        """Capture the OS Level Test creation modal."""
        print("\n[CREATE OS LEVEL TEST]")
        self._goto("/os-level-tests")
        time.sleep(1)
        # Click the New OS Level Test button to open modal
        try:
            self.page.click('button:has-text("New OS Level Test")')
            time.sleep(1)
            self._screenshot("10_create_os_test", "Create OS Level Test",
                "Select servers, load generators, template, cycles, and load profiles.\n"
                "Duration and ramp-up are configurable per load profile.\n"
                "Click 'Create & Start' to create and immediately begin execution.")
        except Exception:
            self._wait_and_screenshot("10_create_os_test", "Create OS Level Test (page)",
                "OS Level Test creation page.")

    def capture_test_dashboard_completed(self):
        """Find a completed test and capture its dashboard."""
        print("\n[TEST DASHBOARD]")
        self._goto("/baseline-tests")
        self.page.wait_for_load_state("networkidle")
        time.sleep(2)

        # Try to find a completed test dashboard link
        try:
            # Look for a dashboard button in the table
            links = self.page.query_selector_all('a[title="Dashboard"]')
            if links:
                # Click the first one
                href = links[0].get_attribute("href")
                if href:
                    self._goto(href)
                    self._wait_and_screenshot("09_test_dashboard", "Test Dashboard - Live View",
                        "The test dashboard shows real-time progress during execution:\n"
                        "- Target Progress: per-server state and current load profile\n"
                        "- Load Profile Progress: calibration and execution status per LP and cycle\n"
                        "- Calibration details: thread count, CPU readings, stability checks")
                    return
        except Exception:
            pass

        print("  No test dashboards available to capture")

    def capture_test_results(self):
        """Find a completed test and capture results page."""
        print("\n[TEST RESULTS]")
        self._goto("/baseline-tests")
        self.page.wait_for_load_state("networkidle")
        time.sleep(2)

        try:
            links = self.page.query_selector_all('a[title="Results"]')
            if links:
                href = links[0].get_attribute("href")
                if href:
                    self._goto(href)
                    self._wait_and_screenshot("10_test_results", "Test Results - Comparison",
                        "After a compare test completes, the results page shows:\n"
                        "- Overall verdict (passed/failed)\n"
                        "- Per-profile comparison results with CPU delta analysis\n"
                        "- Stored profile data: thread counts, stats files, JTL files\n"
                        "- Summary of agent impact on system metrics")
                    return
        except Exception:
            pass

        print("  No test results available to capture")

    def capture_trending(self):
        """Capture trending page."""
        print("\n[TRENDING]")
        self._goto("/trending")
        self._wait_and_screenshot("11_trending", "Trending Analytics",
            "Track performance metrics across multiple test runs over time. "
            "Filter by server, load profile, and date range to identify trends in agent overhead.")

    def generate_word_doc(self, output_path):
        """Generate Word document with screenshots and descriptions."""
        print(f"\n[GENERATING DOCX] {output_path}")

        doc = Document()

        # Title
        title = doc.add_heading("Agent Performance Testing Orchestrator", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("Runbook & User Guide")
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(100, 100, 100)

        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_para.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(150, 150, 150)

        doc.add_page_break()

        # Table of Contents header
        doc.add_heading("Table of Contents", level=1)
        for i, (_, section_title, _) in enumerate(self.screenshots, 1):
            toc_para = doc.add_paragraph()
            toc_para.paragraph_format.space_after = Pt(2)
            run = toc_para.add_run(f"{i}. {section_title}")
            run.font.size = Pt(11)

        doc.add_page_break()

        # Overview section
        doc.add_heading("Overview", level=1)
        doc.add_paragraph(
            "This runbook documents the Agent Performance Testing Orchestrator, "
            "a distributed performance testing framework that measures the CPU overhead "
            "of security agents on servers.\n\n"
            "The system works by:\n"
            "1. Calibrating thread counts to achieve target CPU levels (20-40%, 40-60%, 60-80%)\n"
            "2. Running controlled workloads via JMeter against emulator endpoints\n"
            "3. Collecting CPU/memory statistics during execution\n"
            "4. Comparing results between baseline (no agent) and agent-installed snapshots\n\n"
            "Key components:\n"
            "- Orchestrator: FastAPI web application managing test lifecycle\n"
            "- Emulator: Java Spring Boot application on targets simulating server workloads\n"
            "- JMeter: Load generator sending HTTP requests to emulators\n"
            "- Hypervisor: Proxmox/vSphere for snapshot management and VM operations"
        )

        doc.add_page_break()

        # Workflow section
        doc.add_heading("Test Workflow", level=1)
        doc.add_paragraph(
            "New Baseline Test Flow:\n"
            "  1. Validating - Pre-flight checks (connectivity, snapshots, packages)\n"
            "  2. Deploying Loadgen - Revert loadgen, install JMeter\n"
            "  3. Deploying Calibration - Revert targets, deploy emulator, upload JMX+CSV\n"
            "  4. Calibrating - Find thread count for each load profile's target CPU range\n"
            "  5. Generating - Create deterministic ops sequence CSVs\n"
            "  6. Deploying Testing - Revert targets, deploy emulator for test run\n"
            "  7. Executing - Run JMeter at calibrated thread counts, collect stats\n"
            "  8. Storing - Save results to DB and disk\n\n"
            "Compare Test Flow:\n"
            "  1. Validating\n"
            "  2. Deploying Loadgen\n"
            "  3. Deploying Testing - Uses stored thread counts + CSVs from parent baseline\n"
            "  4. Executing - Runs same workload on new snapshot (with agent)\n"
            "  5. Comparing - Computes CPU delta between baseline and agent snapshots\n"
            "  6. Storing"
        )

        doc.add_page_break()

        # Screenshots sections
        for i, (img_path, section_title, description) in enumerate(self.screenshots, 1):
            doc.add_heading(f"{i}. {section_title}", level=1)

            if description:
                for para_text in description.split("\n"):
                    if para_text.strip():
                        doc.add_paragraph(para_text.strip())

            # Add screenshot
            if os.path.exists(img_path):
                doc.add_paragraph()  # spacing
                try:
                    doc.add_picture(img_path, width=Inches(6.5))
                    last_para = doc.paragraphs[-1]
                    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    doc.add_paragraph(f"[Screenshot: {img_path} - Error: {e}]")

            doc.add_page_break()

        # Appendix: Key Commands
        doc.add_heading("Appendix: Key Commands", level=1)
        commands = [
            ("Start Orchestrator",
             "cd orchestrator\nPYTHONPATH=src python -m uvicorn orchestrator.app:app --host 0.0.0.0 --port 8000"),
            ("Sync Code from GitHub",
             "cd setup\npython sync_setup.py\nfind ../orchestrator -name '*.pyc' -delete"),
            ("Setup JMeter on a Machine",
             "python setup_jmeter.py --disk sdc"),
            ("Retake Snapshots",
             "python retake_snapshots.py 'test name' --sudo-user ak1svc231"),
            ("Validate Calibration",
             "python validate_calibration.py 'test name' --check-running"),
            ("Live Calibration Test",
             "python test_calibration_live.py --target-ip 10.x.x.x --threads 8 --think-ms 100"),
            ("Test Workload (Direct HTTP)",
             "python test_workload.py --target-ip 10.x.x.x --touch-mb 1.0"),
            ("Delete Cache Files",
             "find ../orchestrator -name '*.pyc' -delete\n"
             "find ../orchestrator -name '__pycache__' -type d -exec rm -rf {} + 2>/dev/null"),
        ]
        for title, cmd in commands:
            doc.add_heading(title, level=2)
            doc.add_paragraph(cmd, style="No Spacing")
            doc.add_paragraph()

        # Save
        doc.save(output_path)
        print(f"  [OK] Saved: {output_path}")


def main():
    parser = argparse.ArgumentParser(description="Generate runbook with screenshots")
    parser.add_argument("--url", default="http://localhost:8000",
                        help="Orchestrator base URL")
    parser.add_argument("--user", default="admin", help="Login username")
    parser.add_argument("--password", default="admin", help="Login password")
    parser.add_argument("--output", default=None,
                        help="Output docx path (default: runbook_<date>.docx)")
    parser.add_argument("--screenshot-dir", default="runbook_screenshots",
                        help="Directory for screenshots")
    args = parser.parse_args()

    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    output_path = args.output or f"runbook_{timestamp}.docx"

    gen = RunbookGenerator(args.url, args.user, args.password, args.screenshot_dir)

    print(f"{'='*60}")
    print(f"  RUNBOOK GENERATOR")
    print(f"  URL: {args.url}")
    print(f"  Output: {output_path}")
    print(f"  Screenshots: {args.screenshot_dir}/")
    print(f"{'='*60}")

    with sync_playwright() as p:
        gen.browser = p.chromium.launch(headless=True)
        gen.page = gen.browser.new_page(viewport={"width": 1400, "height": 900})

        try:
            gen.login()
            gen.capture_dashboard()
            gen.capture_servers()
            gen.capture_load_profiles()
            gen.capture_agents()
            gen.capture_agent_sets()
            gen.capture_snapshot_manager()
            gen.capture_tests_page()
            gen.capture_baseline_test_list()
            gen.capture_create_new_baseline()
            gen.capture_test_dashboard_completed()
            gen.capture_test_results()
            gen.capture_trending()
        except Exception as e:
            print(f"\n  [ERROR] Screenshot capture failed: {e}")
            import traceback
            traceback.print_exc()
        finally:
            gen.browser.close()

    # Generate Word doc
    gen.generate_word_doc(output_path)

    print(f"\n{'='*60}")
    print(f"  DONE")
    print(f"  Word document: {output_path}")
    print(f"  Screenshots: {args.screenshot_dir}/")
    print(f"{'='*60}")


if __name__ == "__main__":
    main()
